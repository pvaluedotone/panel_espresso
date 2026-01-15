"""
Export Report Module
Exports comprehensive analysis results to Word documents

This module provides functions to convert formatted text reports
(including regression tables, diagnostics, and statistics) into
publication-ready Word documents.

Author: Saiyidi MAT RONI
Date: December 24, 2025
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from datetime import datetime
from typing import Optional
import os


def create_formatted_document() -> Document:
    """
    Create a new Word document with standard formatting styles
    
    Returns:
    --------
    Document : python-docx Document object with pre-configured styles
    """
    doc = Document()
    
    # Configure default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Create custom styles if they don't exist
    styles = doc.styles
    
    # Heading style for main sections
    try:
        heading_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
    except:
        heading_style = styles['Heading 1']
    
    # Subheading style
    try:
        subheading_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.name = 'Calibri'
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True
        subheading_style.font.color.rgb = RGBColor(0, 102, 204)
    except:
        subheading_style = styles['Heading 2']
    
    # Code/monospace style
    try:
        code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(51, 51, 51)
    except:
        code_style = styles['Normal']
    
    return doc


def parse_table_from_text(text: str) -> Optional[list]:
    """
    Parse ASCII table from text and return as list of rows
    
    Parameters:
    -----------
    text : str
        Text containing ASCII table with | or consistent spacing
    
    Returns:
    --------
    list or None : List of lists representing table rows, or None if no table found
    """
    lines = text.strip().split('\n')
    
    # Look for lines with consistent | separators or spacing
    table_lines = []
    for line in lines:
        # Skip separator lines (----, ═══, etc.)
        if re.match(r'^[\s\-═│┼├┤┬┴╪╫╬╠╣╦╩╬]*$', line):
            continue
        # Check if line looks like table content
        if '|' in line or re.search(r'\s{3,}', line):
            table_lines.append(line)
    
    if not table_lines:
        return None
    
    # Parse rows
    rows = []
    for line in table_lines:
        # Split by | or multiple spaces
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells
        else:
            cells = re.split(r'\s{3,}', line.strip())
        
        if cells:
            rows.append(cells)
    
    return rows if len(rows) > 1 else None


def add_formatted_text(doc: Document, text: str):
    """
    Add formatted text to document, parsing special formatting
    
    Handles:
    - Headers (lines starting with # or enclosed in ══════)
    - Bold text (**text** or __text__)
    - Code blocks (monospace sections)
    - Bullet points (lines starting with • or -)
    - Tables (ASCII tables with consistent structure)
    
    Parameters:
    -----------
    doc : Document
        python-docx Document object
    text : str
        Formatted text to add
    """
    if not text or not text.strip():
        return
    
    sections = text.split('\n\n')
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        # Check for box headers (╔═══╗ style)
        if re.match(r'^[╔═╗╚╝║]+$', section):
            continue
        
        # Check for header in box
        if section.startswith('║') and section.endswith('║'):
            header_text = section.strip('║ ')
            p = doc.add_paragraph(header_text)
            p.style = 'CustomHeading1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # Check for markdown headers
        if section.startswith('#'):
            level = len(re.match(r'^#+', section).group())
            header_text = section.lstrip('#').strip()
            if level == 1:
                p = doc.add_paragraph(header_text)
                p.style = 'CustomHeading1'
            else:
                p = doc.add_paragraph(header_text)
                p.style = 'CustomHeading2'
            continue
        
        # Check for section headers (lines with ─── above/below)
        if '─' in section or '═' in section:
            lines = section.split('\n')
            for line in lines:
                if not re.match(r'^[\s\-═│┼├┤┬┴╪╫╬╠╣╦╩╬]*$', line):
                    p = doc.add_paragraph(line)
                    p.style = 'CustomHeading2'
            continue
        
        # Check for tables
        table_data = parse_table_from_text(section)
        if table_data and len(table_data) > 1:
            # Create table in Word
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        cell = table.rows[i].cells[j]
                        cell.text = cell_data
                        # Make first row bold (header)
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
            doc.add_paragraph()  # Add spacing after table
            continue
        
        # Check for code blocks (lines with consistent indentation)
        if section.startswith('    ') or section.startswith('\t'):
            p = doc.add_paragraph(section)
            p.style = 'CodeBlock'
            continue
        
        # Check for bullet points
        lines = section.split('\n')
        is_list = all(line.strip().startswith(('•', '-', '*', '→', '✓', '✗', '⚠️', '📊', '📚', '📖')) 
                     for line in lines if line.strip())
        
        if is_list:
            for line in lines:
                line = line.strip()
                if line:
                    # Remove bullet marker
                    line = re.sub(r'^[•\-\*→✓✗⚠️📊📚📖]\s*', '', line)
                    p = doc.add_paragraph(line, style='List Bullet')
            continue
        
        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        
        # Parse inline formatting
        parts = re.split(r'(\*\*[^\*]+\*\*|__[^_]+__|`[^`]+`)', section)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('__') and part.endswith('__'):
                # Bold text (alternative)
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                # Code/monospace
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Regular text
                run = p.add_run(part)


def export_report_to_word(
    report_text: str,
    output_filename: Optional[str] = None,
    output_dir: Optional[str] = None
) -> str:
    """
    Export comprehensive report to Word document
    
    Parameters:
    -----------
    report_text : str
        The formatted text report from generate_publication_report()
    output_filename : str, optional
        Custom filename for the output document
        Default: "Panel_Analysis_Report_YYYYMMDD_HHMMSS.docx"
    output_dir : str, optional
        Directory to save the document
        Default: Current working directory
    
    Returns:
    --------
    str : Full path to the created Word document
    """
    # Create document with formatting
    doc = create_formatted_document()
    
    # Add title page
    title = doc.add_paragraph("Panel Data Analysis Report")
    title.style = 'CustomHeading1'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Add page break after title
    doc.add_page_break()
    
    # Add report content
    add_formatted_text(doc, report_text)
    
    # Add footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Panel Data Analysis | Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Determine output path
    if output_filename is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"Panel_Analysis_Report_{timestamp}.docx"
    
    if not output_filename.endswith('.docx'):
        output_filename += '.docx'
    
    if output_dir is None:
        output_dir = os.getcwd()
    
    output_path = os.path.join(output_dir, output_filename)
    
    # Save document
    doc.save(output_path)
    
    return output_path


def export_report_to_word_gradio(report_text: str, output_dir: str = "./", 
                                  stored_results=None, stored_df=None, 
                                  stored_params=None) -> str:
    """
    Gradio-friendly wrapper for export_report_to_word
    Returns a message with the file path
    
    NEW: Optionally includes stepwise models if stored results provided
    
    Parameters:
    -----------
    report_text : str
        The formatted text report
    output_dir : str
        Directory to save the document (default: current directory)
    stored_results : optional
        Stored results object for stepwise models
    stored_df : optional
        Stored dataframe for stepwise models
    stored_params : optional
        Stored parameters for stepwise models
    
    Returns:
    --------
    str : Success message with file path, or error message
    """
    if not report_text or report_text.strip() == "":
        return "❌ Error: No report to export. Please run the analysis first and generate a comprehensive report."
    
    if "No results available" in report_text or "Error" in report_text:
        return "❌ Error: Cannot export empty or error report. Please run a successful analysis first."
    
    try:
        # Check if we should include stepwise models
        include_stepwise = all([stored_results is not None, 
                               stored_df is not None, 
                               stored_params is not None])
        
        # Add stepwise report if available
        full_report = report_text
        if include_stepwise:
            try:
                from stepwise_report import generate_stepwise_report
                stepwise_text = generate_stepwise_report(stored_results, stored_df, stored_params)
                if stepwise_text and "Error" not in stepwise_text:
                    full_report = report_text + "\n\n" + "=" * 80 + "\n\n" + stepwise_text
            except Exception as e:
                print(f"Warning: Could not include stepwise report: {str(e)}")
        
        output_path = export_report_to_word(full_report, output_dir=output_dir)
        abs_path = os.path.abspath(output_path)
        
        stepwise_msg = "\n- **Stepwise regression models** (progressive specifications)" if include_stepwise else ""
        
        return f"""
✅ **Report exported successfully!**

📄 **File location**: `{abs_path}`

**File size**: {os.path.getsize(abs_path) / 1024:.1f} KB

You can now:
- Open the document in Microsoft Word
- Share with collaborators
- Include in your manuscript or presentation

💡 **Tip**: The Word document includes:
- Formatted regression tables
- Descriptive statistics
- Model diagnostics
- Bootstrap results (if applicable){stepwise_msg}
- All tables are editable and publication-ready
"""
    
    except Exception as e:
        return f"❌ **Export failed**: {str(e)}\n\nPlease check that you have write permissions for the output directory."


# ═══════════════════════════════════════════════════════════════════════════
#                    GRADIO INTEGRATION COMPONENTS
# ═══════════════════════════════════════════════════════════════════════════


def create_export_ui():
    """
    Create Gradio UI components for export functionality
    
    Returns:
    --------
    tuple : (export_button, export_status)
        - export_button: Button to trigger export
        - export_status: Textbox to display export status
    """
    with gr.Row():
        export_button = gr.Button(
            "📥 Export to Word Document",
            variant="primary",
            size="lg"
        )
    
    export_status = gr.Textbox(
        label="Export status",
        lines=8,
        interactive=False
    )
    
    return export_button, export_status
